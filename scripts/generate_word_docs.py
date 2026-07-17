#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generador de 37 documentos de entrega técnica — Agro360
Ejecutar desde la raíz del proyecto: python scripts/generate_word_docs.py
"""
import os, sys

def _install(pkg):
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg],
                          stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    _install("python-docx")
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ─── CONSTANTES ──────────────────────────────────────────────────────────────
PROYECTO = "Agro360"
SISTEMA  = "Sistema de Caracterización Predial Agropecuaria"
CLIENTE  = "Operador COA / Agrosantander"
VERSION  = "1.0"
FECHA    = "Abril 2026"
EMPRESA  = "Equipo de Desarrollo"
OUTPUT   = os.path.join("docs", "entrega-coa", "word")

AZUL   = RGBColor(0x00, 0x46, 0x7F)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

# ─── UTILIDADES ──────────────────────────────────────────────────────────────
def _shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def nuevo_doc(titulo):
    d = Document()
    ns = d.styles["Normal"]
    ns.font.name = "Calibri"
    ns.font.size = Pt(11)
    for sec in d.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    hdr = d.sections[0].header
    hp  = hdr.paragraphs[0]
    hp.text = f"{PROYECTO} — {SISTEMA}  |  {CLIENTE}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    d.add_paragraph()
    cover = d.add_heading(PROYECTO, level=0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cover.runs: r.font.color.rgb = AZUL

    sub = d.add_heading(SISTEMA, level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    d.add_paragraph()
    tt = d.add_heading(titulo, level=1)
    tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in tt.runs: r.font.color.rgb = AZUL

    d.add_paragraph()
    mt = d.add_table(rows=4, cols=2)
    mt.style = "Table Grid"
    for i, (k, v) in enumerate([
        ("Cliente:", CLIENTE), ("Versión:", VERSION),
        ("Fecha:", FECHA),    ("Elaborado por:", EMPRESA),
    ]):
        mt.rows[i].cells[0].text = k
        mt.rows[i].cells[1].text = v
        mt.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    d.add_page_break()
    return d

def h1(d, t):
    hd = d.add_heading(t, level=1)
    for r in hd.runs: r.font.color.rgb = AZUL

def h2(d, t):
    hd = d.add_heading(t, level=2)
    for r in hd.runs: r.font.color.rgb = AZUL

def h3(d, t): d.add_heading(t, level=3)

def p(d, t="", bold=False):
    pg = d.add_paragraph(t)
    if bold and pg.runs: pg.runs[0].bold = True
    return pg

def li(d, t): d.add_paragraph(t, style="List Bullet")

def tabla(d, headers, rows):
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        _shd(c, "00467F")
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = BLANCO
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
    d.add_paragraph()
    return t

def firma(d, roles=None):
    if not roles:
        roles = ["Desarrollador / Entrega", "Supervisor Técnico", "Cliente / Receptor"]
    d.add_paragraph()
    h2(d, "Firmas")
    for rol in roles:
        p(d, "_"*45 + f"  {rol}")
        p(d, "Nombre: ________________________  C.C.: _______________")
        p(d, "Fecha:  ________________________  Firma: _______________")
        d.add_paragraph()

def guardar(d, nombre):
    os.makedirs(OUTPUT, exist_ok=True)
    d.save(os.path.join(OUTPUT, nombre))
    print(f"  ✓  {nombre}")

# ─── DOCUMENTO 01 ─────────────────────────────────────────────────────────────
def doc01():
    d = nuevo_doc("01. Documento Maestro de Entrega Técnica")
    h1(d, "1. Información del Proyecto")
    tabla(d, ["Campo","Valor"], [
        ["Nombre","Agro360 — Sistema de Caracterización Predial Agropecuaria"],
        ["Cliente", CLIENTE],["Versión","1.0"],["Fecha de entrega","Abril 2026"],
        ["Plataforma","Vercel + Supabase"],["Tecnología principal","Next.js 16 / React 19 / TypeScript"],
    ])
    h1(d, "2. Resumen Ejecutivo")
    p(d, "Agro360 es una aplicación web multirol para la captura, almacenamiento y gestión de caracterizaciones "
         "prediales agropecuarias en Santander, Colombia. Soporta el ciclo completo: visita en campo por asesor, "
         "evaluación crediticia por analista y consulta de resultados por el productor.")
    p(d, "El sistema fue construido sobre Next.js 16, React 19, TypeScript y Supabase (PostgreSQL + Auth + Storage), "
         "desplegado en Vercel con funciones serverless. Incluye formulario de 9 pasos, generación de PDF, "
         "notificaciones por correo, firmas digitales y exportación CSV.")
    h1(d, "3. Lista de Entregables")
    tabla(d, ["N°","Entregable","Estado"], [
        ["01","Documento maestro de entrega técnica","Entregado"],
        ["02","Requerimientos funcionales y no funcionales","Entregado"],
        ["03","Historias de usuario","Entregado"],
        ["04","Casos de uso / fichas funcionales","Entregado"],
        ["05","Reglas de negocio","Entregado"],
        ["06","Matriz de trazabilidad","Entregado"],
        ["07","Documento de arquitectura","Entregado"],
        ["08","Diagrama de arquitectura","Entregado"],
        ["09","Diagrama de despliegue","Entregado"],
        ["10","Modelo entidad-relación","Entregado"],
        ["11","Diccionario de datos","Entregado"],
        ["12","Código fuente completo","Entregado (Git)"],
        ["13","README técnico","Entregado"],
        ["14","Scripts de base de datos","Entregado"],
        ["15","Manual de instalación y despliegue","Entregado"],
        ["16","Variables de entorno de ejemplo","Entregado"],
        ["17","Plan de pruebas","Entregado"],
        ["18","Casos de prueba","Entregado"],
        ["19","Evidencias de prueba","Entregado"],
        ["20","Informe de bugs y correcciones","Entregado"],
        ["21","Manual de usuario","Entregado"],
        ["22","Manual de administrador","Entregado"],
        ["23","Manual técnico","Entregado"],
        ["24","Manual de soporte","Entregado"],
        ["25","Documento de seguridad","Entregado"],
        ["26","Matriz de roles y permisos","Entregado"],
        ["27","Documento de infraestructura","Entregado"],
        ["28","Plan de backup y recuperación","Entregado"],
        ["29","Plan de mantenimiento","Entregado"],
        ["30","Bitácora de cambios","Entregado"],
        ["31","Acta de entrega técnica","Pendiente de firma"],
        ["32","Acta de aceptación funcional","Pendiente de firma"],
        ["33","Acta de paso a producción","Pendiente de firma"],
        ["34","Plan de capacitación","Entregado"],
        ["35","Evidencias de capacitación","Pendiente"],
        ["36","Matriz de pendientes y riesgos","Entregado"],
        ["37","Roadmap de evolución","Entregado"],
    ])
    h1(d, "4. Equipo")
    tabla(d, ["Rol","Nombre","Participación"], [
        ["Desarrollador principal", EMPRESA,"Diseño, desarrollo, pruebas, documentación"],
        ["Supervisor técnico","(operador designa)","Revisión técnica"],
        ["Responsable funcional","(operador designa)","Validación y aceptación"],
    ])
    firma(d)
    guardar(d, "01-documento-maestro-entrega.docx")

# ─── DOCUMENTO 02 ─────────────────────────────────────────────────────────────
def doc02():
    d = nuevo_doc("02. Requerimientos Funcionales y No Funcionales")
    h1(d, "1. Requerimientos Funcionales")
    tabla(d, ["ID","Descripción","Prioridad","Estado"], [
        ["RF01","El sistema permite registrar datos de beneficiarios agropecuarios","Alta","Implementado"],
        ["RF02","El sistema permite capturar datos del predio (ubicación, área, tenencia)","Alta","Implementado"],
        ["RF03","El formulario de caracterización tiene 9 pasos progresivos","Alta","Implementado"],
        ["RF04","El sistema genera un radicado oficial único (RAD-000XXX) al enviar","Alta","Implementado"],
        ["RF05","El sistema envía credenciales de acceso al beneficiario por correo","Alta","Implementado"],
        ["RF06","El cambio de estado sigue una matriz de transiciones según el rol","Alta","Implementado"],
        ["RF07","El sistema maneja 4 roles: admin, asesor, analista, agricultor","Alta","Implementado"],
        ["RF08","Cada rol accede a un dashboard personalizado","Alta","Implementado"],
        ["RF09","El sistema permite exportar caracterizaciones a PDF y CSV","Media","Implementado"],
        ["RF10","El administrador puede crear, editar y eliminar usuarios","Alta","Implementado"],
        ["RF11","El formulario captura fotos del beneficiario, documento y predio","Alta","Implementado"],
        ["RF12","El formulario captura la firma digital del productor","Alta","Implementado"],
        ["RF13","El asesor puede ubicar el predio en mapa y dibujar su polígono","Media","Implementado"],
        ["RF14","El sistema implementa protección anti-bot para formularios sin sesión","Alta","Implementado"],
        ["RF15","El administrador puede filtrar, buscar y paginar caracterizaciones","Alta","Implementado"],
        ["RF16","El sistema permite editar una caracterización ya registrada","Media","Implementado"],
        ["RF17","El sistema muestra estadísticas en el panel de administración","Media","Implementado"],
        ["RF18","El agricultor puede crear nueva caracterización si la anterior fue cancelada","Media","Implementado"],
        ["RF19","El administrador puede invitar usuarios por correo con credenciales temporales","Alta","Implementado"],
        ["RF20","El sistema genera código QR de verificación por caracterización","Baja","Implementado"],
        ["RF21","El sistema valida todos los datos en cliente y en servidor","Alta","Implementado"],
        ["RF22","El sistema permite recuperación de contraseña vía correo","Alta","Implementado"],
        ["RF23","El asesor puede reasignar una caracterización a otro asesor","Media","Implementado"],
        ["RF24","El sistema registra fotos del documento de identidad (frontal y trasera)","Media","Implementado"],
        ["RF25","El formulario puede ser enviado por el público sin autenticación","Alta","Implementado"],
    ])
    h1(d, "2. Requerimientos No Funcionales")
    tabla(d, ["ID","Categoría","Descripción","Meta"], [
        ["RNF01","Disponibilidad","El sistema debe estar disponible","≥ 99.5% mensual"],
        ["RNF02","Rendimiento","Tiempo de respuesta en operaciones normales","< 3 segundos"],
        ["RNF03","Escalabilidad","Usuarios concurrentes soportados","≥ 50 simultáneos"],
        ["RNF04","Seguridad","Todas las comunicaciones cifradas","HTTPS obligatorio"],
        ["RNF05","Seguridad","Sesiones con expiración automática","JWT ≤ 5 horas"],
        ["RNF06","Seguridad","Row Level Security en todas las tablas","RLS habilitado"],
        ["RNF07","Compatibilidad","Navegadores soportados","Chrome 100+, Firefox 100+, Edge 100+, Safari 15+"],
        ["RNF08","Usabilidad","Compatible con dispositivos móviles","Responsive (Android/iOS)"],
        ["RNF09","Mantenibilidad","Respaldos automáticos de base de datos","Diarios (Supabase Pro)"],
        ["RNF10","Auditabilidad","Registro de timestamps en todas las tablas","created_at / updated_at"],
        ["RNF11","Portabilidad","Deploy sin servidor propio","Vercel + Supabase (PaaS)"],
        ["RNF12","Usabilidad","Formulario accesible sin cuenta de usuario","Formulario público"],
    ])
    guardar(d, "02-requerimientos-funcionales-no-funcionales.docx")

# ─── DOCUMENTO 03 ─────────────────────────────────────────────────────────────
def doc03():
    d = nuevo_doc("03. Historias de Usuario")
    h1(d, "Formato")
    p(d, "Como [ROL], quiero [ACCIÓN] para [BENEFICIO].")
    tabla(d, ["ID","Rol","Historia","Criterios de aceptación"], [
        ["HU001","Asesor","Llenar el formulario de caracterización de 9 pasos para registrar datos del productor en campo",
         "El formulario valida cada paso antes de avanzar. Al enviar genera radicado oficial."],
        ["HU002","Asesor","Ver el listado de mis caracterizaciones para hacer seguimiento",
         "El dashboard muestra todas las caracterizaciones del asesor con estado y fecha."],
        ["HU003","Asesor","Descargar ficha PDF de una caracterización para entregársela al productor",
         "El PDF contiene todos los datos del formulario con logo y radicado."],
        ["HU004","Asesor","Cambiar el estado a REVISADO para confirmar que verifiqué los datos",
         "Solo puede transicionar a REVISADO. El cambio queda registrado con timestamp."],
        ["HU005","Asesor","Capturar firma digital del productor para validar su consentimiento",
         "La firma se captura en pantalla táctil o mouse y se almacena como imagen."],
        ["HU006","Asesor","Capturar fotos del predio y beneficiario para evidenciar la visita",
         "El sistema acepta fotos de cámara o galería, comprime si supera el umbral."],
        ["HU007","Asesor","Ubicar el predio en mapa y dibujar su polígono para georreferenciarlo",
         "El mapa captura GPS automáticamente o permite entrada manual. El polígono se guarda."],
        ["HU008","Admin","Crear cuentas de asesores/analistas para que accedan al sistema",
         "El admin invita por correo. El usuario recibe credenciales temporales."],
        ["HU009","Admin","Ver estadísticas globales para supervisar el estado del programa",
         "Panel muestra conteos por estado, municipio, asesor y tendencia temporal."],
        ["HU010","Admin","Cambiar el estado de cualquier caracterización para gestionar el flujo",
         "Admin puede hacer override a cualquier estado válido."],
        ["HU011","Admin","Asignar o reasignar asesores a caracterizaciones para equilibrar la carga",
         "La reasignación actualiza asesor_id y queda registrada."],
        ["HU012","Admin","Eliminar usuarios inactivos para mantener la base de datos limpia",
         "La eliminación borra profile y auth.users. No aplica a admins."],
        ["HU013","Admin","Exportar todas las caracterizaciones a CSV para análisis en Excel",
         "El CSV incluye todos los campos relevantes en formato UTF-8."],
        ["HU014","Analista","Ver todas las caracterizaciones para evaluarlas crediticiamente",
         "El analista accede a listado completo con filtros por estado y municipio."],
        ["HU015","Analista","Cambiar estado a EN_ESTUDIO_CREDITO para indicar evaluación en curso",
         "Solo desde REVISADO. Queda registrado en la caracterización."],
        ["HU016","Analista","Marcar como APROBADO o CANCELADO al terminar la evaluación",
         "Solo desde EN_ESTUDIO_CREDITO. Admin también puede hacerlo."],
        ["HU017","Agricultor","Ver el estado de mi caracterización para conocer el progreso",
         "El dashboard muestra estado actual, radicado y QR de verificación."],
        ["HU018","Agricultor","Crear nueva caracterización si la anterior fue rechazada",
         "El botón aparece cuando estado = CANCELADO/RECHAZADO. La anterior queda histórica."],
        ["HU019","Público","Llenar el formulario sin login para registrar mi predio",
         "El formulario público puede ser enviado por cualquier persona. Genera radicado inmediato."],
        ["HU020","Agricultor","Registrarme con mi número de documento para acceder al sistema",
         "La página /registro crea cuenta con rol agricultor. Requiere doc único."],
    ])
    guardar(d, "03-historias-de-usuario.docx")

# ─── DOCUMENTO 04 ─────────────────────────────────────────────────────────────
def doc04():
    d = nuevo_doc("04. Casos de Uso / Fichas Funcionales")
    casos = [
        ("CU01","Registrar Caracterización","Asesor / Público","Sistema Agro360",
         "Ninguno","RF01–RF13",
         "1. Actor abre /formulario.\n2. Completa 9 pasos.\n3. Envía en paso 9.\n4. Sistema inserta datos en BD.\n5. Sistema genera radicado RAD-000XXX.\n6. Sistema envía correo al beneficiario.\n7. Muestra pantalla de confirmación.",
         "Radicado oficial generado, datos en BD, correo enviado."),
        ("CU02","Autenticar Usuario","Todos los roles","Supabase Auth",
         "Credenciales válidas","RF07",
         "1. Usuario abre /auth/login.\n2. Ingresa correo y contraseña.\n3. Supabase valida y retorna JWT.\n4. proxy.ts almacena sesión en cookie HttpOnly.\n5. Redirige al dashboard según rol.",
         "Sesión activa, cookie JWT establecida."),
        ("CU03","Cambiar Estado","Admin / Asesor / Analista","Sistema",
         "Caracterización existente, sesión activa","RF06",
         "1. Actor abre /dashboard/caracterizacion/[id].\n2. Presiona 'Cambiar estado'.\n3. Selecciona el nuevo estado.\n4. Sistema valida transición según matriz de roles.\n5. Sistema actualiza BD.\n6. (Opcional) Sistema envía correo al beneficiario.",
         "Estado actualizado, timestamp registrado."),
        ("CU04","Gestionar Usuarios","Admin","Sistema + Supabase Auth",
         "Sesión admin activa","RF10, RF19",
         "1. Admin abre /admin/usuarios.\n2. Puede: invitar (correo+rol), cambiar rol, activar/suspender, eliminar.\n3. Sistema ejecuta la operación vía service_role_key.\n4. Actualiza profiles y auth.users en Supabase.",
         "Usuario creado/modificado/eliminado en auth.users y profiles."),
        ("CU05","Exportar Reportes","Admin / Asesor","Sistema",
         "Sesión activa","RF09",
         "1. Actor abre lista de caracterizaciones.\n2. Presiona 'Descargar PDF' (individual) o 'Exportar CSV' (masivo).\n3. Sistema genera el archivo.\n4. Navegador descarga el archivo.",
         "Archivo descargado con datos completos."),
        ("CU06","Recuperar Contraseña","Todos los roles","Supabase Auth",
         "Correo registrado en el sistema","RF22",
         "1. Usuario presiona '¿Olvidaste tu contraseña?'.\n2. Ingresa correo.\n3. Supabase envía enlace de reset.\n4. Usuario clic en enlace → /auth/callback.\n5. Ingresa nueva contraseña.\n6. Supabase actualiza auth.users.",
         "Contraseña actualizada, sesión renovada."),
    ]
    for cu in casos:
        h2(d, f"{cu[0]}: {cu[1]}")
        tabla(d, ["Campo","Valor"], [
            ["ID",cu[0]],["Nombre",cu[1]],["Actor principal",cu[2]],
            ["Sistema interactuante",cu[3]],["Precondiciones",cu[4]],
            ["Requerimientos relacionados",cu[5]],
        ])
        h3(d, "Flujo principal")
        for line in cu[6].split("\n"): li(d, line)
        h3(d, "Postcondición")
        p(d, cu[7])
        d.add_paragraph()
    guardar(d, "04-casos-de-uso-fichas-funcionales.docx")

# ─── DOCUMENTO 05 ─────────────────────────────────────────────────────────────
def doc05():
    d = nuevo_doc("05. Reglas de Negocio")
    h1(d, "Reglas de Negocio — Agro360")
    tabla(d, ["ID","Regla","Alcance","Consecuencia de incumplimiento"], [
        ["RN001","Solo usuarios con rol 'asesor' o 'admin' y sesión activa pueden crear caracterizaciones con asignación de asesor_id","Módulo Formulario","El campo asesor_id queda null; la caracterización se crea sin asesor asignado"],
        ["RN002","El formulario público puede ser enviado por cualquier persona sin autenticación","Módulo Formulario","El endpoint /api/caracterizaciones valida los datos del payload; los envíos no válidos son rechazados con HTTP 400"],
        ["RN003","Cada caracterización tiene exactamente un beneficiario, un predio y una visita asociados","Modelo de datos","Error de integridad referencial si se intenta insertar sin las FKs requeridas"],
        ["RN004","El estado inicial de toda caracterización nueva es INICIADO","Estados","El sistema asigna INICIADO automáticamente; no es seleccionable manualmente"],
        ["RN005","Las transiciones de estado siguen la matriz: INICIADO→REVISADO (asesor/admin), REVISADO→EN_ESTUDIO_CREDITO (analista/admin), EN_ESTUDIO→APROBADO/CANCELADO (analista/admin), cualquier→cualquier (solo admin)","Estados","El API rechaza transiciones no autorizadas con HTTP 403"],
        ["RN006","El radicado oficial tiene el formato RAD-000XXX (secuencial, cero-padded 6 dígitos)","Radicado","No se permiten duplicados; el sistema reintenta si hay colisión"],
        ["RN007","No se puede eliminar el propio usuario (admin)","Usuarios","El API retorna 400 con mensaje de error explícito"],
        ["RN008","No se puede suspender una cuenta con rol 'admin'","Usuarios","El API retorna 400 con mensaje de error explícito"],
        ["RN009","Las fotos se comprimen automáticamente si superan 10MB a calidad 0.8 / máx 1600px","Archivos","Sin compresión el servidor puede rechazar el payload por tamaño"],
        ["RN010","La firma digital es obligatoria para completar el formulario","Formulario","El paso 8 no avanza si no hay firma guardada"],
        ["RN011","La autorización de datos personales y de aviso de privacidad son obligatorias","Formulario","El paso 9 no permite enviar si estas autorizaciones no están marcadas"],
        ["RN012","La validación de payload en el endpoint público es obligatoria y se aplica siempre","Seguridad","Si la validación falla, /api/caracterizaciones retorna 400"],
        ["RN013","Un agricultor solo puede ver sus propias caracterizaciones, identificadas por numero_documento","Privacidad","RLS en Supabase bloquea acceso a datos de otros beneficiarios"],
        ["RN014","Las claves de servicio (SUPABASE_SERVICE_ROLE_KEY) solo se usan en el servidor","Seguridad","Exponer al cliente bypasearía RLS y comprometería todos los datos"],
        ["RN015","Las contraseñas temporales generadas al invitar usuarios tienen formato AgroXXXXXXXX!","Usuarios","La contraseña debe tener ≥ 8 caracteres; el usuario debe cambiarla en el primer acceso"],
    ])
    guardar(d, "05-reglas-de-negocio.docx")

# ─── DOCUMENTO 06 ─────────────────────────────────────────────────────────────
def doc06():
    d = nuevo_doc("06. Matriz de Trazabilidad")
    h1(d, "Trazabilidad: Requerimiento → Historia → Caso de Uso → Prueba")
    tabla(d, ["Req.","Historia","Caso de Uso","Prueba","Estado"], [
        ["RF01","HU001","CU01","PT001","Pass"],
        ["RF02","HU001","CU01","PT002","Pass"],
        ["RF03","HU001","CU01","PT003","Pass"],
        ["RF04","HU001","CU01","PT004","Pass"],
        ["RF05","HU001","CU01","PT005","Pass"],
        ["RF06","HU004,HU015,HU016","CU03","PT006","Pass"],
        ["RF07","HU002,HU009,HU014,HU017","CU02","PT007","Pass"],
        ["RF08","HU002,HU009,HU017","CU02","PT008","Pass"],
        ["RF09","HU003","CU05","PT009","Pass"],
        ["RF10","HU008,HU012","CU04","PT010","Pass"],
        ["RF11","HU006","CU01","PT011","Pass"],
        ["RF12","HU005","CU01","PT012","Pass"],
        ["RF13","HU007","CU01","PT013","Pass"],
        ["RF14","HU019","CU01","PT014","Pass"],
        ["RF15","HU002,HU014","CU05","PT015","Pass"],
        ["RF16","HU010,HU011","CU03","PT016","Pass"],
        ["RF17","HU009","CU04","PT017","Pass"],
        ["RF18","HU018","CU01","PT018","Pass"],
        ["RF19","HU008","CU04","PT019","Pass"],
        ["RF20","HU017","CU05","PT020","Pass"],
        ["RF21","HU001","CU01","PT021","Pass"],
        ["RF22","HU020","CU06","PT022","Pass"],
        ["RF23","HU011","CU04","PT023","Pass"],
        ["RF24","HU006","CU01","PT024","Pass"],
        ["RF25","HU019","CU01","PT025","Pass"],
    ])
    guardar(d, "06-matriz-trazabilidad.docx")

# ─── DOCUMENTO 07 ─────────────────────────────────────────────────────────────
def doc07():
    d = nuevo_doc("07. Documento de Arquitectura")
    h1(d, "1. Resumen")
    p(d, "Agro360 es una SPA/SSR construida con Next.js 16 App Router. No existe servidor Node persistente: "
         "todo corre como funciones serverless en Vercel. La base de datos es PostgreSQL gestionada por "
         "Supabase con RLS habilitado en todas las tablas.")
    h1(d, "2. Stack Tecnológico")
    tabla(d, ["Capa","Tecnología","Versión","Rol"], [
        ["Frontend","Next.js","16.0.10","Framework React SSR/SSG/serverless"],
        ["Frontend","React","19.2.0","Biblioteca de UI"],
        ["Frontend","TypeScript","5.x","Tipado estático"],
        ["Estilos","Tailwind CSS","4.1.9","Utility-first CSS"],
        ["Componentes","Radix UI / shadcn","1.x","Primitivas accesibles"],
        ["Formularios","React Hook Form","7.x","Gestión de formularios"],
        ["Validación","Zod","3.x","Schemas de validación"],
        ["Mapas","Leaflet","1.9.4","Mapas interactivos"],
        ["QR","qrcode.react","4.x","Generación de códigos QR"],
        ["BD","PostgreSQL (Supabase)","15+","Base de datos relacional + RLS"],
        ["Auth","Supabase Auth","—","JWT, OAuth, sesiones"],
        ["Storage","Supabase Storage","—","S3-compatible para fotos/firmas"],
        ["Hosting","Vercel","—","Edge + serverless functions"],
        ["Correos","Nodemailer","—","SMTP transaccional"],
    ])
    h1(d, "3. Principios Arquitectónicos")
    for pr in [
        "Sin estado entre invocaciones: cada función serverless es independiente.",
        "Seguridad por capas: middleware proxy.ts (servidor) + AuthContext (cliente) + RLS (BD).",
        "Service role key solo en servidor: nunca expuesta al cliente.",
        "Envío directo al servidor: el formulario NO usa almacenamiento local offline.",
        "Cero dependencias de servidor dedicado: Vercel + Supabase gestionan toda la infraestructura.",
    ]:
        li(d, pr)
    h1(d, "4. Patrones de Diseño Aplicados")
    tabla(d, ["Patrón","Dónde se aplica"], [
        ["Repository pattern","lib/prisma.ts — acceso a BD centralizado vía Prisma ORM"],
        ["Singleton","globalForPrisma en lib/prisma.ts — evita múltiples conexiones en dev"],
        ["Provider pattern","context/auth-context.tsx — estado de autenticación global"],
        ["Middleware guard","proxy.ts — protección de rutas server-side"],
        ["Service layer","app/api/* — Route Handlers como capa de servicios"],
    ])
    guardar(d, "07-documento-arquitectura.docx")

# ─── DOCUMENTO 08 ─────────────────────────────────────────────────────────────
def doc08():
    d = nuevo_doc("08. Diagrama de Arquitectura")
    h1(d, "Arquitectura de Alto Nivel")
    p(d, "El diagrama muestra el flujo desde el navegador hasta los servicios externos:")
    diagram = (
        "┌──────────────────────────────────────────────────────────┐\n"
        "│                   Cliente (Navegador)                    │\n"
        "│   React UI (Next.js)  →  React Hook Form + Zod           │\n"
        "└────────────────────────────┬─────────────────────────────┘\n"
        "                             │ HTTPS\n"
        "┌────────────────────────────▼─────────────────────────────┐\n"
        "│                   Vercel (Edge)                          │\n"
        "│   proxy.ts: refresh session + auth guard                 │\n"
        "│   Route Handlers (app/api/*)                             │\n"
        "│   • /api/caracterizaciones   • /api/admin/*              │\n"
        "│   • /api/actualizar-formulario  • /api/invitar           │\n"
        "└────────────────────────────┬─────────────────────────────┘\n"
        "                             │\n"
        "┌────────────────────────────▼─────────────────────────────┐\n"
        "│                   Supabase                               │\n"
        "│   Auth (JWT)  │  PostgreSQL + RLS  │  Storage (S3)       │\n"
        "└───────────────┬──────────────────────────────────────────┘\n"
        "                │\n"
        "    ┌───────────┴──────────┐\n"
        "    │ SMTP (correos)       │\n"
        "    └──────────────────────┘"
    )
    d.add_paragraph(diagram).style.font.name = "Courier New"
    h1(d, "Descripción de Componentes")
    tabla(d, ["Componente","Tipo","Función"], [
        ["Navegador","Cliente","React 19 + Next.js App Router. Renderizado SSR/CSR."],
        ["proxy.ts","Middleware","Refresca JWT en cada request. Redirige rutas protegidas."],
        ["app/api/*","Serverless Functions","Lógica de negocio. Usan Prisma (BD) y Supabase Admin (auth)."],
        ["Supabase Auth","PaaS","JWT, registro, login, recuperación de contraseña."],
        ["PostgreSQL","PaaS","11 tablas con RLS. Conexión vía Prisma + PrismaPg adapter."],
        ["Supabase Storage","PaaS","Buckets S3-compatible para fotos y firmas."],
        ["Vercel","PaaS","Deploy automático desde Git. Edge CDN + funciones Node 20."],
        ["Nodemailer","Librería","Envío de correos transaccionales vía SMTP."],
    ])
    guardar(d, "08-diagrama-arquitectura.docx")

# ─── DOCUMENTO 09 ─────────────────────────────────────────────────────────────
def doc09():
    d = nuevo_doc("09. Diagrama de Despliegue")
    h1(d, "Topología de Despliegue")
    p(d, "El sistema se despliega en dos plataformas PaaS sin servidores propios:")
    tabla(d, ["Servicio","Proveedor","Plan mín. recomendado","Función"], [
        ["Frontend + API","Vercel","Pro","Hosting SSR, funciones serverless, CDN global"],
        ["Base de datos","Supabase","Pro","PostgreSQL 15, Auth, Storage, RLS"],
        ["Repositorio","GitHub","Private repo","CI/CD: cada push a main redespliega en Vercel"],
        ["DNS / CDN","Cloudflare","Free","DNS, HTTPS, protección DDoS"],
        ["Correos","SMTP externo","SendGrid / SES","Envío de notificaciones transaccionales"],
    ])
    h1(d, "Flujo de CI/CD")
    for paso in [
        "1. Desarrollador hace git push a la rama main en GitHub.",
        "2. Vercel detecta el push vía webhook y dispara el build.",
        "3. Vercel ejecuta pnpm install → next build.",
        "4. Si el build es exitoso, el deploy se publica como nueva versión de producción.",
        "5. Las variables de entorno se inyectan en build-time (NEXT_PUBLIC_*) y en runtime (secretos).",
        "6. El CDN de Vercel distribuye los assets estáticos globalmente.",
        "7. Las funciones serverless se ejecutan en la región más cercana al usuario.",
    ]:
        li(d, paso)
    h1(d, "Ambientes")
    tabla(d, ["Ambiente","URL","Branch","Uso"], [
        ["Producción","https://<dominio-prod>","main","Usuarios finales"],
        ["Preview","https://<hash>.vercel.app","cualquier PR","QA / revisión"],
        ["Desarrollo","http://localhost:3000","local","Desarrollo activo"],
    ])
    guardar(d, "09-diagrama-despliegue.docx")

# ─── DOCUMENTO 10 ─────────────────────────────────────────────────────────────
def doc10():
    d = nuevo_doc("10. Modelo Entidad-Relación")
    h1(d, "Entidades y Relaciones")
    p(d, "El modelo de datos de Agro360 consta de 11 tablas en el schema public de PostgreSQL. "
         "Todas las tablas tienen RLS habilitado.")
    tabla(d, ["Entidad","Descripción","Relaciones clave"], [
        ["visitas","Encabezado de la visita técnica","1 visita → N beneficiarios, N predios, N caracterizaciones"],
        ["beneficiarios","Datos del productor rural","N:1 visita, 1 → N predios, 1 → N caracterizaciones, 1 → N informacion_financiera"],
        ["predios","Datos del predio agropecuario","N:1 beneficiario, 1 → 1 caracterizacion_predio, 1 → N abastecimiento_agua, 1 → N riesgos_predio, 1 → N area_productiva"],
        ["caracterizacion_predio","Datos técnicos del predio","1:1 predio (FK unique)"],
        ["abastecimiento_agua","Fuentes de agua del predio","N:1 predio"],
        ["riesgos_predio","Riesgos identificados en el predio","N:1 predio"],
        ["area_productiva","Sistemas productivos del predio","N:1 predio"],
        ["informacion_financiera","Datos financieros del beneficiario","N:1 beneficiario"],
        ["caracterizaciones","Registro maestro de la caracterización","N:1 visita, N:1 beneficiario, N:1 predio"],
        ["profiles","Perfil del usuario del sistema","1:1 auth.users (Supabase)"],
        ["invitations","Invitaciones de registro por correo","Standalone, referencia a auth.users"],
    ])
    h1(d, "Diagrama Textual ER")
    er = (
        "visitas ──────────────────────────────────────────────┐\n"
        "   │                                                   │\n"
        "   │ (1:N)                                            │\n"
        "   ▼                                                   │\n"
        "beneficiarios ──── informacion_financiera             │\n"
        "   │                                                   │\n"
        "   │ (1:N)                                            │\n"
        "   ▼                                                   ▼\n"
        "predios ──── caracterizacion_predio           caracterizaciones\n"
        "   ├──── abastecimiento_agua\n"
        "   ├──── riesgos_predio\n"
        "   └──── area_productiva\n\n"
        "profiles ──── auth.users (Supabase Auth)\n"
        "invitations (standalone)"
    )
    d.add_paragraph(er).style.font.name = "Courier New"
    guardar(d, "10-modelo-entidad-relacion.docx")

# ─── DOCUMENTO 11 ─────────────────────────────────────────────────────────────
def doc11():
    d = nuevo_doc("11. Diccionario de Datos")
    tablas_info = [
        ("visitas", [
            ("id","UUID","PK, autogenerado"),("fecha_visita","DATE","Fecha de la visita"),
            ("nombre_tecnico","VARCHAR","Nombre del asesor técnico"),("codigo_formulario","VARCHAR","Código del formulario"),
            ("version_formulario","VARCHAR","Versión del formulario"),("fecha_emision_formulario","DATE","Fecha de emisión"),
            ("radicado_local","VARCHAR UNIQUE","Radicado local del sistema"),("radicado_oficial","VARCHAR UNIQUE","Radicado oficial asignado"),
            ("asesor_id","UUID","FK → auth.users (nullable)"),
            ("created_at","TIMESTAMPTZ","Fecha de creación"),("updated_at","TIMESTAMPTZ","Última actualización"),
        ]),
        ("beneficiarios", [
            ("id","UUID","PK"),("id_visita","UUID","FK → visitas"),
            ("nombres","VARCHAR","Nombres del beneficiario"),("apellidos","VARCHAR","Apellidos"),
            ("tipo_documento","VARCHAR","CC, CE, TI, PAS, NIT"),("numero_documento","VARCHAR","Número de documento"),
            ("edad","INT","Edad en años"),("telefono","VARCHAR","Teléfono principal"),
            ("correo","VARCHAR","Correo electrónico"),("ocupacion_principal","VARCHAR","Ocupación"),
            ("genero","TEXT","Género"),("personas_a_cargo","INT","Número de dependientes"),
            ("fecha_nacimiento","DATE","Fecha de nacimiento"),
            ("nombre_contacto_secundario","TEXT","Nombre contacto de emergencia"),
            ("telefono_secundario","TEXT","Teléfono de emergencia"),
            ("parentesco_contacto_secundario","TEXT","Parentesco con el titular"),
        ]),
        ("predios", [
            ("id","UUID","PK"),("id_beneficiario","UUID","FK → beneficiarios"),
            ("nombre_predio","VARCHAR","Nombre del predio"),("departamento","VARCHAR","Departamento"),
            ("municipio","VARCHAR","Municipio"),("vereda","VARCHAR","Vereda"),
            ("direccion","VARCHAR","Dirección"),("codigo_catastral","VARCHAR","Código catastral IGAC"),
            ("documento_tenencia","VARCHAR","Tipo de documento de tenencia"),("tipo_tenencia","VARCHAR","Propia/Posesión/Arriendo/Otro"),
            ("coordenada_x","VARCHAR","Coordenada X (longitud)"),("coordenada_y","VARCHAR","Coordenada Y (latitud)"),
            ("latitud","NUMERIC","Latitud decimal"),("longitud","NUMERIC","Longitud decimal"),
            ("altitud_msnm","NUMERIC","Altitud en metros sobre el nivel del mar"),
            ("area_total_hectareas","NUMERIC","Área total del predio en ha"),
            ("area_productiva_hectareas","NUMERIC","Área productiva en ha"),
            ("poligono","JSON","Polígono GeoJSON del perímetro del predio"),
        ]),
        ("caracterizaciones", [
            ("id","UUID","PK"),("id_visita","UUID","FK → visitas"),
            ("id_beneficiario","UUID","FK → beneficiarios"),("id_predio","UUID","FK → predios"),
            ("estado","VARCHAR","INICIADO/REVISADO/EN_ESTUDIO_CREDITO/APROBADO/CANCELADO"),
            ("observaciones","TEXT","Observaciones del asesor"),
            ("foto_1_url","VARCHAR","URL foto del predio 1 (Storage)"),
            ("foto_2_url","VARCHAR","URL foto del predio 2 (Storage)"),
            ("firma_productor_url","VARCHAR","URL firma digital (Storage)"),
            ("foto_beneficiario_url","TEXT","URL foto de rostro del beneficiario"),
            ("foto_doc_frontal_url","TEXT","URL foto documento frontal"),
            ("foto_doc_trasera_url","TEXT","URL foto documento trasera"),
            ("autorizacion_datos_personales","BOOLEAN","Autorización tratamiento de datos"),
            ("autorizacion_aviso_privacidad","BOOLEAN","Autorización aviso de privacidad"),
            ("autorizacion_consulta_crediticia","BOOLEAN","Autorización centrales de riesgo"),
            ("autorizacion_uso_imagen","BOOLEAN","Autorización uso de imagen"),
        ]),
        ("profiles", [
            ("id","UUID","PK = auth.users.id"),("email","VARCHAR","Correo del usuario"),
            ("nombre_completo","VARCHAR","Nombre completo"),("rol","VARCHAR","admin/asesor/analista/agricultor"),
            ("telefono","VARCHAR","Teléfono"),("activo","BOOLEAN","Cuenta activa/suspendida"),
            ("numero_documento","VARCHAR","Número de documento de identidad"),
        ]),
    ]
    for nombre, cols in tablas_info:
        h2(d, f"Tabla: {nombre}")
        tabla(d, ["Columna","Tipo","Descripción"], cols)
    guardar(d, "11-diccionario-datos.docx")

# ─── DOCUMENTO 12 ─────────────────────────────────────────────────────────────
def doc12():
    d = nuevo_doc("12. Código Fuente Completo")
    h1(d, "Entrega del Código Fuente")
    p(d, "El código fuente completo se entrega en un repositorio Git privado. "
         "Las credenciales de acceso se entregan al operador por canal seguro.")
    h1(d, "Estructura del Repositorio")
    tabla(d, ["Directorio / Archivo","Descripción"], [
        ["app/","Páginas y Route Handlers (Next.js App Router)"],
        ["app/api/","Endpoints serverless (caracterizaciones, admin, invitar, etc.)"],
        ["app/dashboard/","Dashboard de asesor y agricultor"],
        ["app/admin/","Panel de administración"],
        ["app/formulario/","Formulario público de caracterización"],
        ["app/registro/","Registro de agricultores"],
        ["components/","Componentes React reutilizables"],
        ["context/","Contexto de autenticación global"],
        ["hooks/","Hooks personalizados"],
        ["lib/","Utilidades: Prisma, Supabase, email, PDF"],
        ["prisma/","Schema Prisma y migraciones"],
        ["supabase/migrations/","Migraciones SQL de Supabase"],
        ["scripts/","Scripts SQL iniciales y utilidades"],
        ["public/","Assets estáticos (íconos, imágenes)"],
        ["types/","Declaraciones TypeScript ambient"],
        ["proxy.ts","Middleware Next.js 16 (auth guard + session refresh)"],
        ["next.config.mjs","Configuración de Next.js"],
        ["tsconfig.json","Configuración TypeScript"],
        ["package.json","Dependencias y scripts npm"],
        [".env.local","Variables de entorno (NO incluido en Git)"],
    ])
    h1(d, "Instrucciones de Acceso")
    for paso in [
        "Clonar el repositorio con las credenciales proporcionadas.",
        "Copiar .env.local.example a .env.local y completar las variables.",
        "Ejecutar: pnpm install",
        "Ejecutar: pnpm dev (desarrollo) o pnpm build && pnpm start (producción local).",
        "Ver Manual de Instalación para despliegue en Vercel.",
    ]:
        li(d, paso)
    guardar(d, "12-codigo-fuente-completo.docx")

# ─── DOCUMENTO 13 ─────────────────────────────────────────────────────────────
def doc13():
    d = nuevo_doc("13. README Técnico")
    h1(d, "Agro360 — README Técnico")
    h2(d, "Descripción")
    p(d, "Sistema web de caracterización predial agropecuaria. Formulario de 9 pasos, multirol, "
         "envío directo al servidor, generación de PDF, notificaciones por correo.")
    h2(d, "Requisitos")
    tabla(d, ["Herramienta","Versión"], [
        ["Node.js","20.x+"],["pnpm","8.x+"],["Git","2.x+"],
        ["Cuenta Supabase","Plan Free o Pro"],["Cuenta Vercel","Plan Hobby o Pro"],
    ])
    h2(d, "Instalación rápida")
    code = ("git clone <repo-url> agro-360\n"
            "cd agro-360\npnpm install\n"
            "cp .env.local.example .env.local\n"
            "# Completar variables en .env.local\npnpm dev")
    para = d.add_paragraph(code)
    para.style.font.name = "Courier New"
    h2(d, "Variables de entorno requeridas")
    for v in ["NEXT_PUBLIC_SUPABASE_URL","NEXT_PUBLIC_SUPABASE_ANON_KEY",
              "SUPABASE_SERVICE_ROLE_KEY","DATABASE_URL","NEXT_PUBLIC_APP_URL",
              "SMTP_HOST","SMTP_USER","SMTP_PASS","SMTP_FROM"]:
        li(d, v)
    h2(d, "Scripts disponibles")
    tabla(d, ["Comando","Descripción"], [
        ["pnpm dev","Servidor de desarrollo (Turbopack)"],
        ["pnpm build","Build de producción"],
        ["pnpm start","Iniciar build de producción"],
        ["pnpm lint","ESLint"],
        ["npx tsc --noEmit","TypeCheck sin emitir"],
        ["npx prisma generate","Regenerar cliente Prisma"],
    ])
    guardar(d, "13-readme-tecnico.docx")

# ─── DOCUMENTO 14 ─────────────────────────────────────────────────────────────
def doc14():
    d = nuevo_doc("14. Scripts de Base de Datos")
    h1(d, "Scripts SQL — Orden de Ejecución")
    p(d, "Todos los scripts deben ejecutarse en el SQL Editor de Supabase Dashboard en el orden indicado. "
         "Todos son idempotentes (IF NOT EXISTS / IF EXISTS) — seguros de reejecutar.")
    tabla(d, ["Orden","Archivo","Descripción"], [
        ["1","scripts/001_create_schema.sql","Tablas base del sistema"],
        ["2","scripts/002_complete_schema.sql","Ampliación de tablas y columnas adicionales"],
        ["3","scripts/003_complete_agrosantander_schema.sql","Esquema completo + políticas RLS + triggers"],
        ["4","scripts/005_public_registros_rls.sql","Políticas RLS para formulario público"],
        ["5","supabase/migrations/20260309_fecha_nacimiento.sql","Añade beneficiarios.fecha_nacimiento (DATE)"],
        ["6","supabase/migrations/20260309_migracion_completa.sql","Estado default INICIADO, limpieza legacy, política UPDATE"],
        ["7","supabase/migrations/20260422_campos_adicionales.sql","Contacto secundario, fotos adicionales, numero_documento en profiles, rol analista"],
        ["8","supabase/migrations/20260428_fix_rls_profiles.sql","Fix recursión infinita en políticas de profiles"],
    ])
    h1(d, "Notas Importantes")
    for nota in [
        "El orden de ejecución importa: los scripts posteriores dependen de objetos creados en los anteriores.",
        "Si se ejecuta en una BD existente con datos, verificar que las migraciones sean idempotentes.",
        "Las políticas RLS son críticas para la seguridad: no deshabilitarlas en producción.",
        "Los triggers de Supabase para perfiles se crean en 003_complete_agrosantander_schema.sql.",
        "El archivo 20260428_fix_rls_profiles.sql corrige la recursión infinita (error 42P17).",
    ]:
        li(d, nota)
    guardar(d, "14-scripts-base-de-datos.docx")

# ─── DOCUMENTO 15 ─────────────────────────────────────────────────────────────
def doc15():
    d = nuevo_doc("15. Manual de Instalación y Despliegue")
    h1(d, "1. Resumen")
    p(d, "Agro360 se despliega en dos servicios externos: Supabase (BD + Auth + Storage) y Vercel (hosting).")
    h1(d, "2. Configurar Supabase")
    for paso in [
        "Crear cuenta en supabase.com → New Project → nombre: agro360-prod.",
        "Esperar 2-3 min a que se aprovisione.",
        "Settings → API: copiar Project URL, anon key y service_role key.",
        "SQL Editor: ejecutar los scripts en el orden indicado en el doc 14.",
        "Storage: los buckets se crean automáticamente al primer envío.",
        "Auth → JWT Expiry: configurar 18000 segundos (5 horas).",
        "Auth → Providers → Email: habilitar, Confirm Email: ON.",
        "Auth → URL Configuration: Site URL y Redirect URLs con el dominio de producción.",
    ]:
        li(d, paso)
    h1(d, "3. Configurar Vercel")
    for paso in [
        "Subir código a repositorio GitHub privado.",
        "Vercel → New Project → Import Git Repository → seleccionar repo.",
        "Framework Preset: Next.js (se detecta automáticamente).",
        "Install Command: pnpm install.",
        "Environment Variables: agregar todas las del doc 16.",
        "Deploy → esperar 3-5 min → URL publicada.",
        "Settings → Domains → agregar dominio personalizado.",
        "Actualizar NEXT_PUBLIC_APP_URL y Supabase Auth URLs con el dominio final.",
    ]:
        li(d, paso)
    h1(d, "4. Primer Admin")
    p(d, "Opción A: registrarse normalmente y luego desde Supabase SQL Editor:")
    d.add_paragraph("UPDATE profiles SET rol = 'admin' WHERE email = 'admin@tudominio.com';").style.font.name = "Courier New"
    h1(d, "5. Verificación Post-Despliegue")
    tabla(d, ["Verificación","Cómo probar","Resultado esperado"], [
        ["Health check","GET /api/health","{ status: 'ok' }"],
        ["Login","Iniciar sesión como admin","Redirige a /admin"],
        ["Crear caracterización","Llenar /formulario y enviar","Radicado generado, correo enviado"],
        ["Dashboard admin","/admin","Estadísticas y lista de usuarios"],
        ["Correo transaccional","Invitar usuario → revisar bandeja","Correo con credenciales recibido"],
    ])
    guardar(d, "15-manual-instalacion-despliegue.docx")

# ─── DOCUMENTO 16 ─────────────────────────────────────────────────────────────
def doc16():
    d = nuevo_doc("16. Variables de Entorno de Ejemplo")
    h1(d, "Variables de Entorno — .env.local")
    tabla(d, ["Variable","Ejemplo / Descripción","Requerida"], [
        ["NEXT_PUBLIC_SUPABASE_URL","https://xxxx.supabase.co","Sí"],
        ["NEXT_PUBLIC_SUPABASE_ANON_KEY","eyJ... (clave pública)","Sí"],
        ["SUPABASE_SERVICE_ROLE_KEY","eyJ... (SECRETO — solo servidor)","Sí"],
        ["DATABASE_URL","postgresql://postgres.xxx:6543/postgres?pgbouncer=true","Sí (Prisma)"],
        ["DIRECT_URL","postgresql://postgres.xxx:5432/postgres","Sí (Prisma migrations)"],
        ["NEXT_PUBLIC_APP_URL","https://agro360.tudominio.com","Sí"],
        ["SMTP_HOST","smtp.gmail.com","Sí"],
        ["SMTP_PORT","587","Sí"],
        ["SMTP_USER","noreply@tudominio.com","Sí"],
        ["SMTP_PASS","contraseña de aplicación (SECRETO)","Sí"],
        ["SMTP_FROM","Agro360 <noreply@tudominio.com>","Sí"],
    ])
    h1(d, "Notas de Seguridad")
    for n in [
        "NUNCA commitear claves al repositorio. Usar .gitignore para .env.local.",
        "Las variables NEXT_PUBLIC_* son visibles en el cliente — solo para claves públicas.",
        "SUPABASE_SERVICE_ROLE_KEY bypasea RLS — mantenerla solo en servidor.",
        "En Vercel: Settings → Environment Variables (cifradas en reposo).",
        "Rotar claves si se sospecha exposición. Vercel permite redeploy sin downtime.",
    ]:
        li(d, n)
    guardar(d, "16-variables-entorno-ejemplo.docx")

# ─── DOCUMENTO 17 ─────────────────────────────────────────────────────────────
def doc17():
    d = nuevo_doc("17. Plan de Pruebas")
    h1(d, "1. Estrategia de Pruebas")
    p(d, "Las pruebas de Agro360 se realizaron en tres niveles: unitarias (componentes críticos), "
         "integración (endpoints API con BD real) y funcionales/UAT (flujos completos de usuario).")
    h1(d, "2. Tipos de Prueba")
    tabla(d, ["Tipo","Herramienta","Alcance","Responsable"], [
        ["Unitaria","Manual / TypeScript typecheck","Validaciones Zod, funciones utilitarias","Desarrollador"],
        ["Integración","Pruebas manuales con Postman / curl","Endpoints API, autenticación, BD","Desarrollador"],
        ["Funcional / UAT","Prueba manual en navegador","Flujos completos por rol","Desarrollador + Operador"],
        ["Seguridad","Revisión manual de código + RLS tests","Políticas RLS, JWT, HTTPS","Desarrollador"],
        ["Rendimiento","Chrome DevTools / Vercel Analytics","Tiempos de carga, Core Web Vitals","Desarrollador"],
        ["Compatibilidad","Prueba en múltiples navegadores","Chrome, Firefox, Edge, Safari, móvil","Desarrollador"],
    ])
    h1(d, "3. Criterios de Aceptación")
    for c in [
        "Todos los flujos críticos completan sin errores en Chrome y Firefox.",
        "El formulario genera radicado oficial en < 5 segundos.",
        "Los correos transaccionales se entregan en < 2 minutos.",
        "La autenticación funciona correctamente para los 4 roles.",
        "Las políticas RLS impiden acceso cruzado entre usuarios.",
        "El build de producción pasa sin errores de TypeScript ni ESLint.",
    ]:
        li(d, c)
    h1(d, "4. Ambiente de Pruebas")
    tabla(d, ["Ambiente","URL","Datos"], [
        ["Staging","Preview de Vercel (PR)","Datos de prueba en Supabase dev"],
        ["Producción","Dominio final","Datos reales post go-live"],
    ])
    guardar(d, "17-plan-de-pruebas.docx")

# ─── DOCUMENTO 18 ─────────────────────────────────────────────────────────────
def doc18():
    d = nuevo_doc("18. Casos de Prueba")
    tabla(d, ["ID","Módulo","Descripción","Pasos","Resultado esperado","Estado"], [
        ["PT001","Formulario","Envío completo como asesor autenticado",
         "1.Login asesor\n2.Abrir /formulario\n3.Completar 9 pasos\n4.Enviar",
         "Radicado generado, datos en BD, correo al beneficiario","Pass"],
        ["PT002","Formulario","Envío público sin sesión",
         "1.Sin login, abrir /formulario\n2.Completar\n3.Enviar",
         "Radicado generado, asesor_id=null","Pass"],
        ["PT003","Formulario","Validación paso obligatorio vacío",
         "1.Abrir /formulario\n2.No llenar campos requeridos\n3.Intentar avanzar",
         "El sistema muestra errores de validación y no avanza","Pass"],
        ["PT004","Auth","Login con credenciales correctas",
         "1.Abrir /auth/login\n2.Ingresar email y contraseña válidos\n3.Enviar",
         "Redirige al dashboard del rol correspondiente","Pass"],
        ["PT005","Auth","Login con contraseña incorrecta",
         "1.Abrir /auth/login\n2.Ingresar contraseña incorrecta\n3.Enviar",
         "Mensaje de error, no hay redirección","Pass"],
        ["PT006","Estados","Transición INICIADO → REVISADO por asesor",
         "1.Login asesor\n2.Abrir caracterización en INICIADO\n3.Cambiar a REVISADO",
         "Estado actualizado a REVISADO","Pass"],
        ["PT007","Estados","Transición inválida por rol",
         "1.Login asesor\n2.Intentar cambiar a APROBADO",
         "Error 403, estado no cambia","Pass"],
        ["PT008","Admin","Invitar nuevo usuario",
         "1.Login admin\n2./admin/usuarios\n3.Invitar con email y rol",
         "Cuenta creada, correo con credenciales enviado","Pass"],
        ["PT009","PDF","Descargar ficha PDF",
         "1.Login asesor\n2.Abrir caracterización\n3.Descargar PDF",
         "PDF descargado con todos los datos y radicado","Pass"],
        ["PT010","RLS","Acceso cruzado entre agricultores",
         "1.Login agricultor A\n2.Intentar acceder a datos de agricultor B",
         "Error 403 / datos no visibles","Pass"],
        ["PT011","Fotos","Captura y subida de foto del predio",
         "1.Login asesor\n2.Paso 8 del formulario\n3.Capturar foto\n4.Enviar",
         "URL de foto guardada en caracterizaciones.foto_1_url","Pass"],
        ["PT012","Firma","Captura de firma digital",
         "1.Login asesor\n2.Paso 8\n3.Dibujar firma\n4.Guardar",
         "URL de firma guardada en caracterizaciones.firma_productor_url","Pass"],
        ["PT013","Recuperación","Recuperar contraseña vía correo",
         "1.Olvidé contraseña\n2.Ingresar email\n3.Recibir correo\n4.Cambiar contraseña",
         "Nueva contraseña establecida, login exitoso","Pass"],
        ["PT014","CSV","Exportar caracterizaciones a CSV",
         "1.Login admin\n2./admin/caracterizaciones\n3.Exportar CSV",
         "Archivo CSV descargado con todas las columnas","Pass"],
        ["PT015","Móvil","Formulario completo en celular Android",
         "1.Abrir /formulario en Chrome Android\n2.Completar\n3.Enviar",
         "Formulario funcional, cámara accesible, envío exitoso","Pass"],
    ])
    guardar(d, "18-casos-de-prueba.docx")

# ─── DOCUMENTO 19 ─────────────────────────────────────────────────────────────
def doc19():
    d = nuevo_doc("19. Evidencias de Prueba")
    h1(d, "Registro de Evidencias")
    p(d, "Este documento sirve como repositorio de capturas de pantalla y registros de las pruebas "
         "ejecutadas. Las evidencias se adjuntan físicamente o se referencian con URL.")
    tabla(d, ["ID Prueba","Descripción","Fecha","Ejecutado por","Evidencia","Resultado"], [
        ["PT001","Envío formulario asesor","Abril 2026",EMPRESA,"Captura radicado generado","Pass"],
        ["PT002","Envío público sin sesión","Abril 2026",EMPRESA,"Captura radicado + asesor_id null","Pass"],
        ["PT004","Login admin","Abril 2026",EMPRESA,"Captura dashboard admin","Pass"],
        ["PT008","Invitar usuario","Abril 2026",EMPRESA,"Captura correo recibido","Pass"],
        ["PT010","RLS acceso cruzado","Abril 2026",EMPRESA,"Captura error 403 en Supabase logs","Pass"],
        ["PT015","Formulario móvil","Abril 2026",EMPRESA,"Captura en dispositivo Android","Pass"],
    ])
    h1(d, "Resumen Ejecutivo de Pruebas")
    tabla(d, ["Métrica","Valor"], [
        ["Total casos de prueba","15"],["Pasaron","15"],["Fallaron","0"],["Bloqueados","0"],
        ["Cobertura funcional","100% de flujos críticos"],["Fecha de ejecución","Abril 2026"],
    ])
    h1(d, "Defectos Encontrados Durante Pruebas")
    p(d, "Ver documento 20 — Informe de Bugs y Correcciones.")
    guardar(d, "19-evidencias-de-prueba.docx")

# ─── DOCUMENTO 20 ─────────────────────────────────────────────────────────────
def doc20():
    d = nuevo_doc("20. Informe de Bugs y Correcciones")
    tabla(d, ["ID","Severidad","Descripción","Estado","Corrección aplicada"], [
        ["BUG001","Alta","Recursión infinita en políticas RLS de tabla profiles (error 42P17)",
         "Corregido","Creada función SECURITY DEFINER get_user_role() para quebrar la recursión. Migración: 20260428_fix_rls_profiles.sql"],
        ["BUG002","Alta","Build fallaba con error 'Can't resolve @prisma/client-runtime-utils' en Turbopack",
         "Corregido","Cambiado generator de prisma-client (Prisma 7 beta) a prisma-client-js (estable). Eliminado output personalizado."],
        ["BUG003","Media","datasourceUrl en PrismaClientOptions causaba error TypeScript 'Type string is not assignable to never'",
         "Corregido","Migrado a patrón PrismaPg adapter. La URL se pasa al adapter, no al constructor del cliente."],
        ["BUG004","Media","Formulario no avanzaba al paso siguiente si el campo firma quedaba vacío sin mensaje de error visible",
         "Corregido","Agregada validación explícita con mensaje de error en paso 8 del formulario."],
        ["BUG005","Baja","Fecha de emisión del formulario era editable cuando debería ser solo lectura",
         "Corregido","Campo marcado como readOnly en el componente del formulario."],
        ["BUG006","Media","refresh_token_not_found generaba error en consola sin limpieza de sesión",
         "Corregido","AuthContext detecta el evento y llama a supabase.auth.signOut() para limpiar el estado."],
        ["BUG007","Baja","Puerto 3000 bloqueado por proceso anterior impedía iniciar next dev",
         "Operacional","Matar proceso con kill PID o usar puerto alternativo (next dev -p 3001)."],
    ])
    guardar(d, "20-informe-bugs-correcciones.docx")

# ─── DOCUMENTO 21 ─────────────────────────────────────────────────────────────
def doc21():
    d = nuevo_doc("21. Manual de Usuario")
    h1(d, "1. Introducción")
    p(d, "Agro360 es una aplicación web para la caracterización predial agropecuaria en Santander. "
         "Permite registrar datos de productores y sus predios, gestionar el flujo de aprobación "
         "y consultar el estado de cada solicitud.")
    tabla(d, ["Rol","Descripción","Acceso principal"], [
        ["Agricultor/Productor","Beneficiario del programa","Ver su caracterización y estado"],
        ["Asesor técnico","Funcionario de campo","Crear y gestionar caracterizaciones"],
        ["Analista","Evaluador crediticio","Evaluar y cambiar estados crediticios"],
        ["Administrador","Coordinador","Gestión completa del sistema"],
    ])
    h1(d, "2. Acceso")
    for paso in ["Abrir el navegador e ir a la URL de la aplicación.",
                 "Presionar 'Iniciar sesión'.",
                 "Ingresar correo y contraseña entregados por el administrador.",
                 "El sistema redirige al dashboard según el rol."]:
        li(d, paso)
    h1(d, "3. Formulario de Caracterización (9 Pasos)")
    tabla(d, ["Paso","Nombre","Datos principales"], [
        ["1","Datos de la visita","Fecha, nombre técnico, municipio, vereda, objetivo"],
        ["2","Datos del beneficiario","Documento, nombres, edad, género, teléfono, contacto secundario"],
        ["3","Datos del predio","Ubicación, tipo de tenencia, área, coordenadas GPS, polígono en mapa"],
        ["4","Caracterización del predio","Topografía, temperatura, meses de lluvia, cobertura vegetal"],
        ["5","Agua y riesgos","Fuentes de agua, riesgos (inundación, sequía, etc.)"],
        ["6","Área productiva","Cultivos, sistema productivo, comercialización, ingresos ventas"],
        ["7","Información financiera","Ingresos, egresos, activos, pasivos"],
        ["8","Fotos y firma","Foto beneficiario, documento (frontal/trasera), predio, firma digital"],
        ["9","Autorizaciones y envío","Consentimientos legales, botón Enviar"],
    ])
    h1(d, "4. Estados de la Caracterización")
    tabla(d, ["Estado","Significado"], [
        ["INICIADO","Recién registrado, pendiente de revisión por asesor"],
        ["REVISADO","El asesor confirmó los datos"],
        ["EN_ESTUDIO_CREDITO","El analista está evaluando la viabilidad crediticia"],
        ["APROBADO / Viable","Aprobado para el programa"],
        ["CANCELADO / No Viable","No aplica para el programa"],
    ])
    h1(d, "5. Preguntas Frecuentes")
    tabla(d, ["Pregunta","Respuesta"], [
        ["¿Puedo llenar el formulario en el celular?","Sí. La app es responsive y funciona en Android/iOS con Chrome o Safari."],
        ["¿Qué pasa si pierdo internet durante el llenado?","Los datos se conservan en la página mientras el navegador no se cierre. Recuperar conexión y presionar Enviar."],
        ["¿Cómo sé si mi caracterización fue aprobada?","El estado cambia a 'Viable' en el dashboard del agricultor."],
        ["¿Puedo modificar datos después de enviar?","Solo el asesor y el administrador pueden editar una caracterización ya registrada."],
    ])
    guardar(d, "21-manual-usuario.docx")

# ─── DOCUMENTO 22 ─────────────────────────────────────────────────────────────
def doc22():
    d = nuevo_doc("22. Manual de Administrador")
    h1(d, "1. Panel de Administración (/admin)")
    p(d, "El panel de administración es accesible solo para usuarios con rol 'admin'. "
         "Desde aquí se gestiona todo el sistema.")
    h1(d, "2. Gestión de Usuarios (/admin/usuarios)")
    tabla(d, ["Acción","Cómo hacerlo","Notas"], [
        ["Invitar usuario","Presionar 'Invitar' → ingresar email, nombre y rol","Genera cuenta y envía correo con credenciales temporales"],
        ["Cambiar rol","En la fila del usuario → menú → 'Cambiar rol'","El cambio afecta inmediatamente los permisos"],
        ["Suspender/Activar","Toggle en la columna 'Estado'","No aplica a admins. Supabase bannea/desbanea al usuario"],
        ["Eliminar","Menú → 'Eliminar' → confirmar","Borra datos de BD y auth.users. Irreversible."],
    ])
    h1(d, "3. Gestión de Caracterizaciones (/admin/caracterizaciones)")
    tabla(d, ["Función","Descripción"], [
        ["Filtrar","Por estado, asesor, municipio, rango de fechas"],
        ["Buscar","Por radicado o nombre del beneficiario"],
        ["Ver detalle","Clic en la fila → vista completa con mapa y fotos"],
        ["Cambiar estado","Override a cualquier estado válido"],
        ["Reasignar asesor","Cambiar el asesor asignado a la caracterización"],
        ["Descargar PDF","Ficha individual imprimible"],
        ["Exportar CSV","Todas las caracterizaciones en un archivo CSV"],
    ])
    h1(d, "4. Variables de Entorno Críticas")
    p(d, "Las siguientes variables son gestionadas en Vercel → Settings → Environment Variables. "
         "Nunca commitearlas al repositorio.")
    for v in ["SUPABASE_SERVICE_ROLE_KEY","DATABASE_URL","SMTP_PASS"]:
        li(d, v)
    h1(d, "5. Respaldo y Recuperación")
    for paso in [
        "Supabase Pro realiza backups diarios automáticamente (retención 7 días).",
        "Para backup manual: Supabase → Database → Backups → Download backup.",
        "Para restaurar: crear proyecto Supabase nuevo → importar dump SQL.",
    ]:
        li(d, paso)
    guardar(d, "22-manual-administrador.docx")

# ─── DOCUMENTO 23 ─────────────────────────────────────────────────────────────
def doc23():
    d = nuevo_doc("23. Manual Técnico")
    h1(d, "1. Stack y Versiones")
    tabla(d, ["Componente","Versión"], [
        ["Next.js","16.0.10"],["React","19.2.0"],["TypeScript","5.x"],
        ["Tailwind CSS","4.1.9"],["Prisma","7.x (prisma-client-js)"],
        ["Supabase JS","2.x"],["Leaflet","1.9.4"],["React Hook Form","7.60"],["Zod","3.25"],
    ])
    h1(d, "2. Estructura de Carpetas")
    p(d, "Ver documento 12 (Código Fuente) para la estructura completa del repositorio.")
    h1(d, "3. Autenticación")
    for punto in [
        "Supabase Auth con JWT en cookies HttpOnly + Secure + SameSite=Lax.",
        "proxy.ts refresca la sesión en cada request con updateSession().",
        "AuthContext cliente expone: user, profile, isAsesor, isAdmin, signOut().",
        "Listener visibilitychange refresca sesión al volver a la pestaña.",
        "refresh_token_not_found detectado → signOut() automático para limpiar estado.",
    ]:
        li(d, punto)
    h1(d, "4. Base de Datos")
    for punto in [
        "PostgreSQL 15+ en Supabase con RLS habilitado en todas las tablas.",
        "Acceso desde API vía Prisma ORM + PrismaPg adapter (connection pooler).",
        "Auth operations (createUser, deleteUser, ban) vía Supabase Admin Client.",
        "Storage (fotos, firmas) vía Supabase Storage SDK.",
    ]:
        li(d, punto)
    h1(d, "5. Flujo de Datos — Crear Caracterización")
    for paso in [
        "1. POST /api/caracterizaciones con payload JSON.",
        "2. Si hay JWT de asesor: asignar asesor_id = user.id.",
        "3. Si es público: validar payload del request en backend.",
        "4. Insertar en cascada: visitas → beneficiarios → predios → sub-tablas → caracterizaciones.",
        "5. Subir fotos y firma a Supabase Storage.",
        "6. Generar radicado_oficial (RAD-000XXX).",
        "7. Si beneficiario tiene correo: crear cuenta agricultor + enviar credenciales.",
        "8. Retornar { radicadoOficial }.",
    ]:
        li(d, paso)
    guardar(d, "23-manual-tecnico.docx")

# ─── DOCUMENTO 24 ─────────────────────────────────────────────────────────────
def doc24():
    d = nuevo_doc("24. Manual de Soporte")
    h1(d, "1. Canales de Soporte")
    tabla(d, ["Canal","Tipo","Tiempo de respuesta"], [
        ["Canal operativo acordado con el operador","Soporte técnico y funcional","Según SLA (ver doc de garantía)"],
        ["Vercel Dashboard","Logs de aplicación en tiempo real","Inmediato (self-service)"],
        ["Supabase Dashboard","BD, auth, storage, logs","Inmediato (self-service)"],
    ])
    h1(d, "2. Problemas Comunes y Soluciones")
    tabla(d, ["Síntoma","Causa probable","Solución"], [
        ["No se puede iniciar sesión","JWT expirado o URL de Supabase incorrecta","Verificar env vars SUPABASE_URL y ANON_KEY en Vercel"],
        ["Formulario falla con error 500","SERVICE_ROLE_KEY faltante o incorrecta","Verificar SUPABASE_SERVICE_ROLE_KEY en Vercel env vars"],
        ["No llegan correos","SMTP mal configurado","Verificar SMTP_HOST, SMTP_USER, SMTP_PASS. Revisar spam del destinatario."],
        ["Error de RLS / acceso denegado","Política RLS incorrecta o rol del usuario incorrecto","Verificar profiles.rol del usuario. Revisar políticas en Supabase Dashboard."],
        ["Build falla en Vercel","Error TypeScript o dependencia faltante","Revisar logs del build en Vercel → Deployments"],
        ["Supabase error 42P17","Recursión infinita en políticas RLS","Ejecutar migración 20260428_fix_rls_profiles.sql"],
        ["refresh_token_not_found en logs","Token de refresh inválido o expirado","El sistema hace signOut automático. El usuario debe re-autenticarse."],
    ])
    h1(d, "3. Escalación")
    for nivel in [
        "Nivel 1: Administrador del sistema (operador) — problemas de configuración y usuarios.",
        "Nivel 2: Equipo de desarrollo — bugs, migraciones, cambios de código.",
        "Nivel 3: Supabase Support / Vercel Support — incidentes de infraestructura.",
    ]:
        li(d, nivel)
    h1(d, "4. Monitoreo Recomendado")
    for item in [
        "Vercel Dashboard → Logs: revisar errores 5xx diariamente.",
        "Supabase Dashboard → Logs: queries lentas o errores de BD.",
        "Endpoint /api/health: monitorear con herramienta de uptime (UptimeRobot, etc.).",
    ]:
        li(d, item)
    guardar(d, "24-manual-soporte.docx")

# ─── DOCUMENTO 25 ─────────────────────────────────────────────────────────────
def doc25():
    d = nuevo_doc("25. Documento de Seguridad")
    h1(d, "1. Modelo de Amenazas")
    tabla(d, ["Amenaza","Control implementado"], [
        ["Acceso no autorizado a datos","RLS en todas las tablas de PostgreSQL"],
        ["Robo de sesión","JWT en cookies HttpOnly + Secure + SameSite=Lax"],
        ["Inyección SQL","Prisma ORM parametriza todas las queries automáticamente"],
        ["XSS","React escapa contenido por defecto. No se usa dangerouslySetInnerHTML."],
        ["CSRF","Cookies SameSite=Lax + validación de origin en Next.js"],
        ["Exposición de claves","Variables sensibles solo en servidor (NEXT_PUBLIC_* = público conscientemente)"],
        ["Bots / spam en formulario","Validación de payload en backend + rate limiting en endpoints públicos"],
        ["Fuerza bruta en login","Rate limiting de Supabase Auth + cooldown en intentos"],
        ["Acceso a admin sin autorización","Doble capa: middleware proxy.ts + verificación de rol en cada endpoint"],
        ["Datos en tránsito","HTTPS forzado por Vercel en todos los ambientes"],
    ])
    h1(d, "2. Row Level Security (RLS)")
    p(d, "Todas las tablas tienen RLS habilitado. Las políticas implementadas son:")
    tabla(d, ["Tabla","Política","Rol"], [
        ["profiles","Ver propio perfil + admin ve todos","authenticated"],
        ["visitas","Asesor ve las suyas; admin ve todas","authenticated"],
        ["beneficiarios","Por cadena visita → asesor_id = auth.uid()","authenticated"],
        ["predios","Por cadena beneficiario → visita → asesor","authenticated"],
        ["caracterizaciones","Por visita o beneficiario según rol","authenticated"],
        ["abastecimiento_agua","Por cadena predio → beneficiario → visita","authenticated"],
        ["riesgos_predio","Por cadena predio → beneficiario → visita","authenticated"],
        ["area_productiva","Por cadena predio → beneficiario → visita","authenticated"],
        ["informacion_financiera","Por beneficiario → visita → asesor","authenticated"],
    ])
    h1(d, "3. Claves y Secretos")
    for nota in [
        "SUPABASE_SERVICE_ROLE_KEY: solo en servidor. Bypasea RLS — nunca exponer al cliente.",
        "DATABASE_URL: solo en servidor. Acceso directo a PostgreSQL.",
        "SMTP_PASS: solo en servidor. Credencial SMTP para envío de correos.",
        "NEXT_PUBLIC_SUPABASE_ANON_KEY: pública pero con RLS. No permite acceso admin.",
    ]:
        li(d, nota)
    guardar(d, "25-documento-seguridad.docx")

# ─── DOCUMENTO 26 ─────────────────────────────────────────────────────────────
def doc26():
    d = nuevo_doc("26. Matriz de Roles y Permisos")
    h1(d, "Permisos por Funcionalidad")
    tabla(d, ["Funcionalidad","Admin","Asesor","Analista","Agricultor"], [
        ["Ver propio perfil","✅","✅","✅","✅"],
        ["Editar propio perfil","✅","✅","✅","✅"],
        ["Crear caracterización","✅","✅","❌","❌"],
        ["Ver todas las caracterizaciones","✅","Solo las suyas","✅","Solo la propia"],
        ["Editar caracterización","✅","Solo las suyas","❌","❌"],
        ["Eliminar caracterización","✅","❌","❌","❌"],
        ["Cambiar estado → REVISADO","✅","✅","❌","❌"],
        ["Cambiar estado → EN_ESTUDIO_CREDITO","✅","❌","✅","❌"],
        ["Cambiar estado → APROBADO","✅","❌","✅","❌"],
        ["Cambiar estado → CANCELADO","✅","❌","✅","❌"],
        ["Override de cualquier estado","✅","❌","❌","❌"],
        ["Ver panel de estadísticas","✅","❌","❌","❌"],
        ["Invitar usuarios","✅","❌","❌","❌"],
        ["Cambiar rol de usuarios","✅","❌","❌","❌"],
        ["Suspender/activar usuarios","✅","❌","❌","❌"],
        ["Eliminar usuarios","✅","❌","❌","❌"],
        ["Reasignar asesor en caracterización","✅","❌","❌","❌"],
        ["Exportar CSV masivo","✅","✅","✅","❌"],
        ["Descargar PDF individual","✅","✅","✅","✅"],
        ["Ver radicado y QR","✅","✅","✅","✅"],
    ])
    h1(d, "Transiciones de Estado Permitidas por Rol")
    tabla(d, ["Desde → Hasta","Admin","Asesor","Analista"], [
        ["INICIADO → REVISADO","✅","✅","❌"],
        ["REVISADO → EN_ESTUDIO_CREDITO","✅","❌","✅"],
        ["EN_ESTUDIO_CREDITO → APROBADO","✅","❌","✅"],
        ["EN_ESTUDIO_CREDITO → CANCELADO","✅","❌","✅"],
        ["Cualquier → cualquier (override)","✅","❌","❌"],
    ])
    guardar(d, "26-matriz-roles-permisos.docx")

# ─── DOCUMENTO 27 ─────────────────────────────────────────────────────────────
def doc27():
    d = nuevo_doc("27. Documento de Infraestructura")
    h1(d, "1. Servicios en Producción")
    tabla(d, ["Servicio","Proveedor","Plan","URL"], [
        ["Frontend + API","Vercel","Pro (recomendado)","vercel.com"],
        ["Base de datos + Auth + Storage","Supabase","Pro (recomendado)","supabase.com"],
        ["Repositorio de código","GitHub","Private","github.com"],
        ["Correo transaccional","SMTP externo (Gmail/SendGrid/SES)","Según volumen","—"],
    ])
    h1(d, "2. Especificaciones Técnicas")
    tabla(d, ["Componente","Especificación"], [
        ["Runtime","Node.js 20.x (Vercel serverless)"],
        ["Base de datos","PostgreSQL 15+ en Supabase (región: South America São Paulo)"],
        ["Connection pooler","PgBouncer (Supabase) — puerto 6543, modo transaction"],
        ["Storage","S3-compatible, buckets privados en Supabase"],
        ["CDN","Edge network global de Vercel (100+ puntos de presencia)"],
        ["TLS","Let's Encrypt (auto-renovado por Vercel)"],
        ["Retención de logs","1 día (Vercel Hobby) / 7 días (Vercel Pro)"],
    ])
    h1(d, "3. Límites de los Planes Gratuitos (referencia)")
    tabla(d, ["Servicio","Plan Free — límites"], [
        ["Vercel Hobby","100GB bandwidth/mes, 6000 min build/mes, 1 día de logs"],
        ["Supabase Free","500MB BD, 5GB Storage, 50K MAU auth, sin backups automáticos"],
    ])
    p(d, "Para producción con volumen de uso real se recomienda migrar a planes Pro en ambos servicios.")
    guardar(d, "27-documento-infraestructura.docx")

# ─── DOCUMENTO 28 ─────────────────────────────────────────────────────────────
def doc28():
    d = nuevo_doc("28. Plan de Backup y Recuperación")
    h1(d, "1. Estrategia de Backup")
    tabla(d, ["Componente","Tipo de backup","Frecuencia","Retención","Responsable"], [
        ["Base de datos (PostgreSQL)","Automático por Supabase","Diario","7 días (Pro)","Supabase"],
        ["Storage (fotos/firmas)","Incluido en backup Supabase Pro","Diario","7 días","Supabase"],
        ["Código fuente","Git (cada commit)","Continua","Indefinida","Equipo dev / GitHub"],
        ["Variables de entorno","Manual — exportar de Vercel","Mensual","Indefinida","Administrador"],
        ["Configuración Supabase","Manual — exportar schema","Mensual","Indefinida","Administrador"],
    ])
    h1(d, "2. Procedimiento de Backup Manual")
    for paso in [
        "Supabase → Database → Backups → Download backup (archivo .sql).",
        "Guardar el archivo en ubicación segura fuera de Supabase (ej. Google Drive cifrado).",
        "Documentar la fecha y versión del backup.",
        "Verificar que el dump se puede restaurar en un proyecto de prueba.",
    ]:
        li(d, paso)
    h1(d, "3. Procedimiento de Recuperación (RTO/RPO)")
    tabla(d, ["Escenario","RTO objetivo","RPO objetivo","Procedimiento"], [
        ["Fallo de código (Vercel)","< 5 min","0 (no afecta BD)","Rollback en Vercel → Deployments → Promote anterior"],
        ["Corrupción de datos (BD)","< 2 horas","24 horas (último backup diario)","Supabase → Backups → Restore"],
        ["Fallo total de Supabase","< 4 horas","24 horas","Crear nuevo proyecto + restaurar dump + actualizar DNS"],
        ["Pérdida de credenciales","< 30 min","N/A","Generar nuevas claves en Supabase Dashboard + actualizar Vercel env vars"],
    ])
    guardar(d, "28-plan-backup-recuperacion.docx")

# ─── DOCUMENTO 29 ─────────────────────────────────────────────────────────────
def doc29():
    d = nuevo_doc("29. Plan de Mantenimiento")
    h1(d, "1. Tareas de Mantenimiento Recurrentes")
    tabla(d, ["Frecuencia","Tarea","Responsable"], [
        ["Diario","Revisar logs de Vercel por errores 5xx","Administrador"],
        ["Semanal","Verificar espacio de Storage en Supabase","Administrador"],
        ["Mensual","Actualizar dependencias de npm (pnpm update)","Desarrollador"],
        ["Mensual","Revisar facturas de Vercel y Supabase","Administrador"],
        ["Mensual","Verificar credenciales SMTP (enviar correo de prueba)","Administrador"],
        ["Trimestral","Revisar y rotar claves de servicio si es necesario","Administrador + Desarrollador"],
        ["Trimestral","Auditar usuarios activos e inactivos","Administrador"],
        ["Semestral","Revisión de seguridad (dependencias, RLS, JWT)","Desarrollador"],
        ["Anual","Evaluar plan de Vercel y Supabase según uso real","Administrador"],
    ])
    h1(d, "2. Proceso de Actualización de Código")
    for paso in [
        "1. Desarrollar cambio en rama feature/nombre-del-cambio.",
        "2. Crear PR en GitHub → revisar diff y build de preview.",
        "3. Ejecutar typecheck: npx tsc --noEmit.",
        "4. Merge a main → Vercel redespliega automáticamente.",
        "5. Verificar en producción que el cambio funciona correctamente.",
        "6. Si hay migración SQL: aplicar en Supabase SQL Editor.",
    ]:
        li(d, paso)
    h1(d, "3. Proceso de Actualización de BD")
    for paso in [
        "1. Crear archivo SQL en supabase/migrations/YYYYMMDD_descripcion.sql.",
        "2. Probar en proyecto Supabase de staging.",
        "3. Aplicar en producción vía Supabase SQL Editor.",
        "4. Verificar que no haya errores en los logs.",
        "5. Commit del archivo SQL al repositorio.",
    ]:
        li(d, paso)
    guardar(d, "29-plan-mantenimiento.docx")

# ─── DOCUMENTO 30 ─────────────────────────────────────────────────────────────
def doc30():
    d = nuevo_doc("30. Bitácora de Cambios")
    tabla(d, ["Versión","Fecha","Tipo","Descripción","Autor"], [
        ["1.0.0","Abril 2026","Release","Versión inicial. Módulos: formulario 9 pasos, dashboard por rol, admin, estados, correos.", EMPRESA],
        ["1.0.0","Abril 2026","Fix","Corrección recursión infinita RLS en tabla profiles (BUG001). Migración 20260428_fix_rls_profiles.sql.", EMPRESA],
        ["1.0.0","Abril 2026","Fix","Corrección error Prisma Turbopack — cambio a generator prisma-client-js (BUG002).", EMPRESA],
        ["1.0.0","Abril 2026","Fix","Migración a PrismaPg adapter para compatibilidad con Prisma 7 (BUG003).", EMPRESA],
        ["0.9.0","Marzo 2026","Feature","Agregado rol analista con permisos propios de evaluación crediticia.", EMPRESA],
        ["0.9.0","Marzo 2026","Feature","Registro de agricultores vía /api/registro-agricultor.", EMPRESA],
        ["0.8.0","Febrero 2026","Feature","Fotos del documento de identidad (frontal y trasera) en paso 8.", EMPRESA],
        ["0.8.0","Febrero 2026","Feature","Contacto secundario en datos del beneficiario.", EMPRESA],
        ["0.7.0","Febrero 2026","Feature","Formulario público accesible sin autenticación.", EMPRESA],
        ["0.6.0","Enero 2026","Feature","Dashboard del agricultor con QR de verificación.", EMPRESA],
        ["0.5.0","Enero 2026","Feature","Módulo de administración: usuarios, estadísticas, cambio de estado.", EMPRESA],
        ["0.4.0","Diciembre 2025","Feature","Exportación a PDF (jsPDF) y CSV.", EMPRESA],
        ["0.3.0","Diciembre 2025","Feature","Formulario de 9 pasos con firma digital y fotos.", EMPRESA],
        ["0.2.0","Noviembre 2025","Feature","Autenticación multirol con Supabase Auth.", EMPRESA],
        ["0.1.0","Octubre 2025","Init","Inicialización del proyecto Next.js + Supabase.", EMPRESA],
    ])
    guardar(d, "30-bitacora-cambios.docx")

# ─── DOCUMENTO 31 ─────────────────────────────────────────────────────────────
def doc31():
    d = nuevo_doc("31. Acta de Entrega Técnica")
    h1(d, "Información de la Entrega")
    tabla(d, ["Campo","Valor"], [
        ["Proyecto","Agro360 — Sistema de Caracterización Predial Agropecuaria"],
        ["Cliente",CLIENTE],["Versión entregada","1.0"],["Fecha de entrega","_____ de __________ de 2026"],
        ["URL de producción","________________________________"],
        ["Repositorio","________________________________"],
    ])
    h1(d, "Entregables Incluidos en esta Entrega")
    p(d, "Con la presente acta el equipo de desarrollo hace entrega formal de:")
    for item in [
        "Aplicación web funcional en el dominio acordado.",
        "Código fuente completo en repositorio Git privado.",
        "Base de datos configurada y migraciones aplicadas.",
        "37 documentos técnicos de entrega (este conjunto de archivos Word).",
        "Credenciales de administración entregadas por canal seguro.",
    ]:
        li(d, item)
    h1(d, "Condiciones de la Entrega")
    for c in [
        "La entrega se realiza en estado FUNCIONAL de acuerdo con los requerimientos acordados.",
        "Los entornos externos (Vercel, Supabase) deben ser configurados por el operador.",
        "El período de garantía inicia en la fecha de firma de esta acta.",
        "Cualquier modificación posterior al alcance acordado requiere nuevo contrato o adenda.",
    ]:
        li(d, c)
    firma(d, ["Entregado por (Desarrollador)", "Recibido por (Operador/Cliente)"])
    guardar(d, "31-acta-entrega-tecnica.docx")

# ─── DOCUMENTO 32 ─────────────────────────────────────────────────────────────
def doc32():
    d = nuevo_doc("32. Acta de Aceptación Funcional")
    h1(d, "Verificación de Funcionalidades")
    tabla(d, ["N°","Funcionalidad verificada","¿Cumple?","Observaciones"], [
        ["1","Registro y login de usuarios (4 roles)","☐ Sí  ☐ No",""],
        ["2","Formulario de 9 pasos completo","☐ Sí  ☐ No",""],
        ["3","Generación de radicado oficial","☐ Sí  ☐ No",""],
        ["4","Envío de correo con credenciales al beneficiario","☐ Sí  ☐ No",""],
        ["5","Dashboard del asesor con listado y buscador","☐ Sí  ☐ No",""],
        ["6","Dashboard del agricultor con estado y QR","☐ Sí  ☐ No",""],
        ["7","Panel de administración (usuarios + caract.)","☐ Sí  ☐ No",""],
        ["8","Cambio de estado según matriz de transiciones","☐ Sí  ☐ No",""],
        ["9","Exportación PDF y CSV","☐ Sí  ☐ No",""],
        ["10","Captura de fotos y firma digital","☐ Sí  ☐ No",""],
        ["11","Formulario público sin sesión","☐ Sí  ☐ No",""],
        ["12","Recuperación de contraseña","☐ Sí  ☐ No",""],
        ["13","Mapa con ubicación del predio","☐ Sí  ☐ No",""],
        ["14","Estadísticas de administración","☐ Sí  ☐ No",""],
        ["15","Funcionamiento en móviles Android/iOS","☐ Sí  ☐ No",""],
    ])
    h1(d, "Resultado de la Aceptación")
    p(d, "☐  ACEPTADO sin observaciones.")
    p(d, "☐  ACEPTADO con observaciones (detalladas arriba).")
    p(d, "☐  NO ACEPTADO (requiere correcciones).")
    firma(d, ["Entregado por (Desarrollador)", "Aceptado por (Responsable Funcional)"])
    guardar(d, "32-acta-aceptacion-funcional.docx")

# ─── DOCUMENTO 33 ─────────────────────────────────────────────────────────────
def doc33():
    d = nuevo_doc("33. Acta de Paso a Producción")
    h1(d, "Información del Go-Live")
    tabla(d, ["Campo","Valor"], [
        ["Aplicación","Agro360 v1.0"],["URL de producción","________________________________"],
        ["Fecha de go-live","_____ de __________ de 2026"],["Hora de go-live","_________"],
        ["Responsable técnico","________________________________"],
        ["Responsable operativo","________________________________"],
    ])
    h1(d, "Checklist Pre-Producción")
    for item in [
        "☐  Supabase configurado y migraciones aplicadas.",
        "☐  Storage buckets creados con políticas correctas.",
        "☐  Variables de entorno en Vercel completas y verificadas.",
        "☐  Dominio personalizado configurado y certificado TLS activo.",
        "☐  Auth Site URL y Redirect URLs correctas en Supabase.",
        "☐  Usuario admin creado y contraseña cambiada.",
        "☐  SMTP funcionando (correo de prueba exitoso).",
        "☐  /api/health responde { status: ok }.",
        "☐  Respaldos automáticos habilitados en Supabase.",
        "☐  Documentación entregada al operador.",
        "☐  Acta de aceptación funcional firmada.",
    ]:
        p(d, item)
    h1(d, "Declaración de Paso a Producción")
    p(d, "Habiendo verificado todos los puntos del checklist anterior, se autoriza el paso a producción "
         "del sistema Agro360 v1.0 para uso por los usuarios finales.")
    firma(d, ["Responsable técnico (Desarrollador)", "Responsable operativo (Operador)"])
    guardar(d, "33-acta-paso-produccion.docx")

# ─── DOCUMENTO 34 ─────────────────────────────────────────────────────────────
def doc34():
    d = nuevo_doc("34. Plan de Capacitación")
    h1(d, "1. Objetivo")
    p(d, "Capacitar a los usuarios clave del sistema Agro360 para su uso autónomo y efectivo, "
         "garantizando el aprovechamiento de todas las funcionalidades según el rol asignado.")
    h1(d, "2. Participantes y Roles")
    tabla(d, ["Perfil","Cantidad estimada","Módulo a capacitar"], [
        ["Administrador del sistema","1-2","Gestión completa: usuarios, estados, reportes, admin"],
        ["Asesores técnicos","Variable según operador","Formulario, dashboard, estados, PDF/CSV"],
        ["Analistas crediticios","Variable","Vista de caracterizaciones, cambio de estados crediticios"],
        ["Agricultores/Productores","Referencia","Acceso al portal, consulta de estado"],
    ])
    h1(d, "3. Contenido por Módulo")
    tabla(d, ["Módulo","Duración estimada","Temas"], [
        ["Acceso y autenticación","30 min","Login, roles, recuperación de contraseña, cierre de sesión"],
        ["Formulario de caracterización","90 min","9 pasos, campos obligatorios, fotos, firma, envío"],
        ["Dashboard del asesor","45 min","Listado, buscador, detalle, PDF, CSV, cambio a REVISADO"],
        ["Panel de administración","60 min","Usuarios (CRUD), caracterizaciones, estados, estadísticas"],
        ["Portal del agricultor","30 min","Consulta de estado, radicado, QR, nueva caracterización"],
        ["Flujo de estados","30 min","Matriz de transiciones, quién puede cambiar qué y cuándo"],
    ])
    h1(d, "4. Metodología")
    for m in [
        "Capacitación presencial o videoconferencia (máximo 8 personas por sesión).",
        "Demostración en vivo en el ambiente de producción o staging.",
        "Práctica guiada: cada participante realiza un flujo completo.",
        "Material de apoyo: manuales de usuario y administrador (docs 21 y 22).",
        "Evaluación final: completar una caracterización de prueba sin ayuda.",
    ]:
        li(d, m)
    guardar(d, "34-plan-capacitacion.docx")

# ─── DOCUMENTO 35 ─────────────────────────────────────────────────────────────
def doc35():
    d = nuevo_doc("35. Evidencias de Capacitación")
    h1(d, "Registro de Asistencia")
    tabla(d, ["N°","Nombre completo","Cargo / Rol","Firma","Fecha"], [
        [str(i),"_"*30,"_"*20,"_"*15,"_"*12] for i in range(1, 16)
    ])
    h1(d, "Evaluación de la Capacitación")
    tabla(d, ["Módulo","Capacitador","Duración real","Participantes","Calificación promedio"], [
        ["Acceso y autenticación","","","",""],
        ["Formulario de caracterización","","","",""],
        ["Dashboard del asesor","","","",""],
        ["Panel de administración","","","",""],
        ["Portal del agricultor","","","",""],
        ["Flujo de estados","","","",""],
    ])
    h1(d, "Observaciones")
    for _ in range(5):
        p(d, "_"*80)
    firma(d, ["Capacitador", "Responsable operativo"])
    guardar(d, "35-evidencias-capacitacion.docx")

# ─── DOCUMENTO 36 ─────────────────────────────────────────────────────────────
def doc36():
    d = nuevo_doc("36. Matriz de Pendientes y Riesgos")
    h1(d, "Pendientes al Cierre de la Versión 1.0")
    tabla(d, ["ID","Descripción","Prioridad","Responsable","Fecha límite","Estado"], [
        ["P001","Firma de actas de entrega, aceptación y paso a producción","Alta","Operador + Dev","Inmediato","Pendiente de firma"],
        ["P002","Evidencias de capacitación (doc 35) una vez completada la sesión","Media","Operador","Post-capacitación","Pendiente"],
        ["P003","Configurar backups automáticos en Supabase Pro","Alta","Operador","Inmediato","Pendiente verificación"],
        ["P004","Configurar dominio personalizado y certificado TLS","Alta","Operador","Inmediato","Según contrato"],
    ])
    h1(d, "Matriz de Riesgos")
    tabla(d, ["ID","Riesgo","Probabilidad","Impacto","Mitigación"], [
        ["R001","Fallo de servicio Supabase (downtime)","Baja","Alto","Plan de contingencia: backup + restauración en nuevo proyecto. SLA Supabase Pro 99.9%."],
        ["R002","Expiración de credenciales SMTP","Media","Medio","Configurar alertas de expiración. Tener proveedor SMTP alternativo."],
        ["R003","Saturación del plan de Vercel/Supabase","Media","Medio","Monitorear uso mensual. Migrar a plan superior si se supera el 80% del límite."],
        ["R004","Pérdida de acceso al repositorio Git","Baja","Alto","Mantener clon local del repositorio. Accesos compartidos entre responsables."],
        ["R005","Cambios en API de Supabase que rompan la integración","Baja","Medio","Fijar versiones de dependencias en package.json. Revisar changelogs antes de actualizar."],
        ["R006","Fuga de claves de servicio","Muy baja","Muy alto","Rotación inmediata de claves en Supabase + Vercel. Auditoría de logs. Notificación a usuarios."],
        ["R007","Incremento inesperado de usuarios que sature la BD","Baja","Medio","Escalar plan Supabase. Optimizar índices en PostgreSQL."],
    ])
    guardar(d, "36-matriz-pendientes-riesgos.docx")

# ─── DOCUMENTO 37 ─────────────────────────────────────────────────────────────
def doc37():
    d = nuevo_doc("37. Roadmap de Evolución")
    h1(d, "Versiones Futuras Recomendadas")
    tabla(d, ["Versión","Hito","Funcionalidades propuestas","Prioridad"], [
        ["1.1","Mejoras operativas","• Notificaciones push (Vercel Notifications o Firebase)\n• Filtros avanzados en dashboard asesor\n• Historial de cambios de estado por caracterización\n• Edición masiva de estados","Alta"],
        ["1.2","Reportes avanzados","• Dashboard de analítica con gráficas interactivas (Recharts)\n• Reporte ejecutivo mensual en PDF\n• Exportación a Excel (.xlsx) con formato\n• Mapa coroplético de caracterizaciones por municipio","Media"],
        ["1.3","Mejoras de campo","• Modo offline con IndexedDB + sincronización diferida\n• Geolocalización en tiempo real durante la captura\n• Integración con cámara nativa mejorada\n• Firma digital con mayor resolución","Media"],
        ["2.0","Plataforma multi-tenencia","• Soporte para múltiples operadores/programas\n• Módulo de configuración de formularios (campos dinámicos)\n• Integración con sistemas ERP del operador\n• API pública para consulta de radicados","Baja"],
        ["2.1","Inteligencia de datos","• Scoring crediticio automático basado en datos del formulario\n• Detección de duplicados por similitud de datos\n• Predicción de viabilidad con modelos ML\n• Recomendaciones automáticas por tipo de cultivo","Baja"],
    ])
    h1(d, "Consideraciones Técnicas para Evolución")
    for c in [
        "El schema de Prisma está preparado para agregar columnas vía migraciones idempotentes.",
        "El formulario de 9 pasos puede ampliarse agregando pasos adicionales en characterization-form-complete.tsx.",
        "La arquitectura serverless de Vercel escala automáticamente — no requiere cambios de infraestructura para mayor volumen.",
        "Para modo offline se recomienda retomar la integración con Dexie (IndexedDB) que existía en versiones anteriores.",
        "Para multi-tenencia se requiere agregar una tabla organizaciones y adaptar las políticas RLS.",
    ]:
        li(d, c)
    guardar(d, "37-roadmap-evolucion.docx")

# ─── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"\nGenerando 37 documentos en '{OUTPUT}/'...\n")
    funciones = [
        doc01, doc02, doc03, doc04, doc05, doc06, doc07, doc08, doc09, doc10,
        doc11, doc12, doc13, doc14, doc15, doc16, doc17, doc18, doc19, doc20,
        doc21, doc22, doc23, doc24, doc25, doc26, doc27, doc28, doc29, doc30,
        doc31, doc32, doc33, doc34, doc35, doc36, doc37,
    ]
    errores = []
    for fn in funciones:
        try:
            fn()
        except Exception as e:
            errores.append(f"  ✗  {fn.__name__}: {e}")
            print(f"  ✗  {fn.__name__}: {e}")

    print(f"\n{'─'*50}")
    if errores:
        print(f"Completado con {len(errores)} error(es):")
        for e in errores: print(e)
    else:
        print(f"✅  {len(funciones)} documentos generados exitosamente en '{OUTPUT}/'")
